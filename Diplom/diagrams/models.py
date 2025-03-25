from django.db import models
from docx import Document as DocxDocument
from accounts.models import Users

class Document(models.Model):
    name = models.CharField(max_length=255)  # Название документа
    file = models.FileField(upload_to='uploaded_docs/')  # Файл
    content = models.TextField(blank=True, null=True)  # Текст из файла

    def save(self, *args, **kwargs):
        """Обновляем content только при загрузке нового файла"""
        if self.file and not self.pk:  # Если создаётся новый объект
            self.content = self.read_docx(self.file.path)
        super().save(*args, **kwargs)

    def read_docx(self, file_path):
        """Функция для чтения содержимого .docx файла"""
        try:
            doc = DocxDocument(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            return f"Ошибка при чтении файла: {str(e)}"

    def __str__(self):
        return self.name


class BPMNFile(models.Model):
    name = models.CharField(max_length=255, default="Новая диаграмма")
    xml_data = models.TextField(blank=True)  # Храним XML-данные диаграммы
    created_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return self.name  # Возвращаем название диаграммы как строку

class BPMNDiagrams(models.Model):
    name = models.CharField(max_length=255, default="Новая диаграмма")
    xml_data = models.TextField()  # Основное поле для хранения XML
    created_at = models.DateTimeField(auto_now_add=True)
    username = models.CharField(max_length=255, blank=True, null=True)  # Новое поле для имени пользователя

    def __str__(self):
        return self.name

class ChatMessage(models.Model):
    user = models.ForeignKey(Users, on_delete=models.CASCADE)  # Связь с пользователем
    message = models.TextField()  # Текст сообщения
    timestamp = models.DateTimeField(auto_now_add=True)  # Время отправки

    def __str__(self):
        return f"{self.user.Username}: {self.message}"

